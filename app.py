import streamlit as st
import pandas as pd
import requests
from io import BytesIO
from datetime import datetime
from docx import Document
from docx.shared import Inches
import matplotlib.pyplot as plt
import numpy as np

# ----------------------------
# Configuração da página e estilo moderno
# ----------------------------
st.set_page_config(page_title="CrowdStrike Dashboard", layout="wide", initial_sidebar_state="collapsed")

st.markdown("""
<style>
    @import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700&display=swap');
    
    .main {
        background-color: #0E1117;
        color: #FFFFFF;
        font-family: 'Inter', sans-serif;
    }
    
    .stApp {
        background: linear-gradient(135deg, #0E1117 0%, #1E1E2E 100%);
    }
    
    h1, h2, h3, h4, h5, h6 {
        color: #FFFFFF !important;
        font-weight: 600 !important;
        font-family: 'Inter', sans-serif !important;
    }
    
    .css-1d391kg, .css-12oz5g7 {
        background-color: #0E1117;
    }
    
    .kpi-card {
        background: linear-gradient(135deg, #1E1E2E 0%, #2D2D44 100%);
        padding: 1.5rem;
        border-radius: 12px;
        border: 1px solid #3D3D5C;
        box-shadow: 0 4px 12px rgba(0,0,0,0.1);
    }
    
    .toggle-container {
        display: flex;
        align-items: center;
        gap: 10px;
        margin: 5px 0;
    }
    
    .toggle-switch {
        position: relative;
        display: inline-block;
        width: 50px;
        height: 24px;
    }
    
    .toggle-switch input {
        opacity: 0;
        width: 0;
        height: 0;
    }
    
    .toggle-slider {
        position: absolute;
        cursor: pointer;
        top: 0;
        left: 0;
        right: 0;
        bottom: 0;
        background-color: #3D3D5C;
        transition: .4s;
        border-radius: 24px;
    }
    
    .toggle-slider:before {
        position: absolute;
        content: "";
        height: 16px;
        width: 16px;
        left: 4px;
        bottom: 4px;
        background-color: white;
        transition: .4s;
        border-radius: 50%;
    }
    
    input:checked + .toggle-slider {
        background-color: #D32F2F;
    }
    
    input:checked + .toggle-slider:before {
        transform: translateX(26px);
    }
    
    .stButton>button {
        background: linear-gradient(135deg, #D32F2F 0%, #B71C1C 100%);
        color: white;
        border: none;
        border-radius: 8px;
        padding: 0.5rem 1.5rem;
        font-weight: 500;
        transition: all 0.3s ease;
    }
    
    .stButton>button:hover {
        background: linear-gradient(135deg, #B71C1C 0%, #9A0007 100%);
        transform: translateY(-2px);
        box-shadow: 0 4px 12px rgba(211, 47, 47, 0.3);
    }
    
    .section-header {
        background: linear-gradient(135deg, #1E1E2E 0%, #2D2D44 100%);
        padding: 1rem 1.5rem;
        border-radius: 10px;
        margin: 1rem 0;
        border-left: 4px solid #D32F2F;
    }
</style>
""", unsafe_allow_html=True)

st.title("🛡️ CrowdStrike Security Dashboard")
st.markdown("---")

# ----------------------------
# Carregar tenants do secrets
# ----------------------------
if "tenants" not in st.secrets:
    st.error("Nenhum tenant configurado no secrets.toml.")
    st.stop()

tenants = st.secrets["tenants"]
tenant_labels = {k: tenants[k]["company_name"] for k in tenants.keys()}
selected_company = st.selectbox("Selecione o Tenant", list(tenant_labels.values()))
selected_key = [k for k, v in tenant_labels.items() if v == selected_company][0]
tenant_cfg = tenants[selected_key]

# ----------------------------
# Funções de API
# ----------------------------
def get_token(cfg):
    url = f"{cfg['base_url']}/oauth2/token"
    data = {"client_id": cfg["client_id"], "client_secret": cfg["client_secret"]}
    headers = {"Content-Type": "application/x-www-form-urlencoded"}
    response = requests.post(url, data=data, headers=headers)
    if response.status_code not in [200, 201]:
        st.error(f"Erro ao obter token ({response.status_code}): {response.text}")
        return None
    return response.json().get("access_token")

def get_all_host_ids(token, cfg):
    url = f"{cfg['base_url']}/devices/queries/devices/v1"
    headers = {"Authorization": f"Bearer {token}"}
    all_ids = []
    offset = 0
    limit = 500
    while True:
        params = {"offset": offset, "limit": limit}
        resp = requests.get(url, headers=headers, params=params)
        if resp.status_code != 200:
            st.error(f"Erro ao buscar IDs: {resp.text}")
            break
        ids_batch = resp.json().get("resources", [])
        if not ids_batch:
            break
        all_ids.extend(ids_batch)
        offset += limit
    return all_ids

def get_hosts_details(token, cfg, ids):
    url = f"{cfg['base_url']}/devices/entities/devices/v2"
    headers = {"Authorization": f"Bearer {token}"}
    all_hosts = []
    for i in range(0, len(ids), 500):
        batch = ids[i:i+500]
        resp = requests.post(url, headers=headers, json={"ids": batch})
        if resp.status_code != 200:
            st.error(f"Erro ao buscar detalhes: {resp.text}")
            continue
        all_hosts.extend(resp.json().get("resources", []))
    if not all_hosts:
        return pd.DataFrame()
    df = pd.json_normalize(all_hosts)
    return df

# ----------------------------
# Filtros avançados com toggle minimalista
# ----------------------------
def aplicar_filtros(df):
    st.markdown('<div class="section-header">🎛️ Filtros Avançados</div>', unsafe_allow_html=True)
    
    col1, col2, col3 = st.columns(3)
    
    filtros = {
        "Sistema Operacional": "os_version",
        "Plataforma": "platform_name", 
        "Versão do Sensor": "agent_version",
        "Política Aplicada": "policy_applied"
    }
    
    toggle_filtros = {
        "Uninstall Protection": "tamper_protection_enabled",
        "RFM": "rfm_enabled", 
        "Duplicada": "is_duplicate"
    }
    
    with col1:
        for label, col in list(filtros.items())[:2]:
            if col in df.columns:
                df[col] = df[col].fillna("Desconhecido")
                valores = ["Todos"] + sorted(df[col].astype(str).unique().tolist())
                escolha = st.selectbox(label, valores, key=f"select_{col}")
                if escolha != "Todos":
                    df = df[df[col].astype(str) == escolha]
    
    with col2:
        for label, col in list(filtros.items())[2:]:
            if col in df.columns:
                df[col] = df[col].fillna("Desconhecido")
                valores = ["Todos"] + sorted(df[col].astype(str).unique().tolist())
                escolha = st.selectbox(label, valores, key=f"select_{col}")
                if escolha != "Todos":
                    df = df[df[col].astype(str) == escolha]
    
    with col3:
        for label, col in toggle_filtros.items():
            if col in df.columns:
                st.markdown(f'<div class="toggle-container"><span>{label}</span>', unsafe_allow_html=True)
                toggle_state = st.radio(label, ["Todos", "Sim", "Não"], horizontal=True, key=f"toggle_{col}", label_visibility="collapsed")
                st.markdown('</div>', unsafe_allow_html=True)
                
                if toggle_state == "Sim":
                    df = df[df[col] == True]
                elif toggle_state == "Não":
                    df = df[df[col] == False]

    return df

# ----------------------------
# Funções para gráficos menores e agrupamento
# ----------------------------
def criar_grafico_barras(df, coluna, titulo, limite_outros=0.05):
    """Cria gráfico de barras com agrupamento de valores pequenos em 'Outros'"""
    if coluna not in df.columns:
        return None
        
    counts = df[coluna].value_counts()
    total = len(df)
    
    # Agrupar valores pequenos em "Outros"
    threshold = total * limite_outros
    principais = counts[counts >= threshold]
    outros = counts[counts < threshold]
    
    if len(outros) > 0:
        principais["Outros"] = outros.sum()
    
    fig, ax = plt.subplots(figsize=(4, 3))
    cores = plt.cm.Reds(np.linspace(0.4, 0.8, len(principais)))
    principais.plot(kind='bar', color=cores, ax=ax)
    
    ax.set_title(titulo, fontsize=11, fontweight='bold', color='white')
    ax.set_facecolor('#0E1117')
    fig.patch.set_facecolor('#0E1117')
    ax.tick_params(colors='white', labelsize=8)
    ax.spines['bottom'].set_color('#3D3D5C')
    ax.spines['top'].set_color('#3D3D5C') 
    ax.spines['right'].set_color('#3D3D5C')
    ax.spines['left'].set_color('#3D3D5C')
    ax.yaxis.label.set_color('white')
    ax.xaxis.label.set_color('white')
    
    plt.xticks(rotation=45, ha='right')
    plt.tight_layout()
    return fig

def criar_grafico_pizza(df, coluna, titulo, limite_outros=0.03):
    """Cria gráfico de pizza com agrupamento de valores pequenos"""
    if coluna not in df.columns:
        return None
        
    counts = df[coluna].value_counts()
    total = len(df)
    
    # Agrupar valores pequenos em "Outros"
    threshold = total * limite_outros
    principais = counts[counts >= threshold]
    outros = counts[counts < threshold]
    
    if len(outros) > 0:
        principais = principais.copy()
        principais["Outros"] = outros.sum()
    
    fig, ax = plt.subplots(figsize=(4, 3))
    cores = plt.cm.Reds(np.linspace(0.4, 0.9, len(principais)))
    principais.plot(kind='pie', ax=ax, colors=cores, autopct='%1.1f%%', startangle=90)
    
    ax.set_title(titulo, fontsize=11, fontweight='bold', color='white')
    ax.set_facecolor('#0E1117')
    fig.patch.set_facecolor('#0E1117')
    ax.tick_params(colors='white', labelsize=8)
    
    plt.tight_layout()
    return fig

# ----------------------------
# Export Word com KPIs e gráficos
# ----------------------------
def export_word(df, tenant_name):
    doc = Document()
    
    # Cabeçalho
    header = doc.sections[0].header
    header_para = header.paragraphs[0]
    header_para.text = f"Relatório Executivo - {tenant_name}"
    header_para.alignment = 1
    
    doc.add_heading(f'Relatório Executivo - {tenant_name}', 0)
    doc.add_paragraph(f'Data: {datetime.now().strftime("%d/%m/%Y %H:%M")}')
    doc.add_paragraph(f'Total Hosts: {len(df)}')
    
    # KPIs
    doc.add_heading("Principais Indicadores", level=1)
    kpi_table = doc.add_table(rows=2, cols=4)
    kpi_table.style = 'Light Grid Accent 1'
    
    # Cabeçalho da tabela
    kpi_headers = ["Sistemas Operacionais", "Versões do Sensor", "RFM Ativo", "Proteção Anti-Desinstalação"]
    for i, header in enumerate(kpi_headers):
        kpi_table.rows[0].cells[i].text = header
    
    # Valores dos KPIs
    kpi_values = [
        df['os_version'].nunique() if 'os_version' in df else 0,
        df['agent_version'].nunique() if 'agent_version' in df else 0,
        df['rfm_enabled'].sum() if 'rfm_enabled' in df else 0,
        df['tamper_protection_enabled'].sum() if 'tamper_protection_enabled' in df else 0
    ]
    
    for i, value in enumerate(kpi_values):
        kpi_table.rows[1].cells[i].text = str(value)
    
    # Gráficos para Word
    doc.add_heading("Distribuições Principais", level=1)
    
    # Gráfico de versões do sensor
    if 'agent_version' in df.columns:
        fig = criar_grafico_barras(df, 'agent_version', 'Versões do Sensor')
        if fig:
            fig_path = "agent_version_plot.png"
            fig.savefig(fig_path, dpi=300, bbox_inches='tight', facecolor='#0E1117')
            plt.close(fig)
            doc.add_picture(fig_path, width=Inches(6))
    
    # Gráfico de sistemas operacionais
    if 'os_version' in df.columns:
        fig = criar_grafico_pizza(df, 'os_version', 'Sistemas Operacionais')
        if fig:
            fig_path = "os_version_plot.png"
            fig.savefig(fig_path, dpi=300, bbox_inches='tight', facecolor='#0E1117')
            plt.close(fig)
            doc.add_picture(fig_path, width=Inches(6))
    
    # Salvar documento em buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return buffer

# ----------------------------
# Dashboard principal
# ----------------------------
st.markdown('<div class="section-header">🔍 Dashboard CrowdStrike</div>', unsafe_allow_html=True)

if st.button("🚀 Buscar Hosts do Tenant", use_container_width=True):
    with st.spinner("Autenticando..."):
        token = get_token(tenant_cfg)
    if not token:
        st.stop()

    with st.spinner("Buscando IDs dos hosts..."):
        ids = get_all_host_ids(token, tenant_cfg)
    if not ids:
        st.warning("Nenhum host encontrado.")
        st.stop()

    with st.spinner("Buscando detalhes dos hosts..."):
        df = get_hosts_details(token, tenant_cfg, ids)
    if df.empty:
        st.warning("Nenhum host retornado.")
        st.stop()

    st.success(f"✅ {len(df)} hosts carregados com sucesso!")

    # Aplicar filtros avançados
    df_filtrado = aplicar_filtros(df)

    # KPIs
    st.markdown('<div class="section-header">📊 KPIs Principais</div>', unsafe_allow_html=True)
    kpi_cols = st.columns(4)
    
    with kpi_cols[0]:
        st.markdown('<div class="kpi-card">', unsafe_allow_html=True)
        st.metric("Total Hosts", len(df_filtrado))
        st.markdown('</div>', unsafe_allow_html=True)
    
    with kpi_cols[1]:
        st.markdown('<div class="kpi-card">', unsafe_allow_html=True)
        st.metric("Sistemas Operacionais", df_filtrado["os_version"].nunique() if "os_version" in df_filtrado else 0)
        st.markdown('</div>', unsafe_allow_html=True)
    
    with kpi_cols[2]:
        st.markdown('<div class="kpi-card">', unsafe_allow_html=True)
        st.metric("Versões do Sensor", df_filtrado["agent_version"].nunique() if "agent_version" in df_filtrado else 0)
        st.markdown('</div>', unsafe_allow_html=True)
    
    with kpi_cols[3]:
        st.markdown('<div class="kpi-card">', unsafe_allow_html=True)
        st.metric("RFM Ativo", f"{df_filtrado['rfm_enabled'].sum() if 'rfm_enabled' in df_filtrado else 0} / {len(df_filtrado)}")
        st.markdown('</div>', unsafe_allow_html=True)

    st.markdown("---")

    # Gráficos
    st.markdown('<div class="section-header">📈 Visualizações</div>', unsafe_allow_html=True)
    
    col1, col2 = st.columns(2)
    
    with col1:
        if 'agent_version' in df_filtrado.columns:
            fig = criar_grafico_barras(df_filtrado, 'agent_version', 'Distribuição de Versões do Sensor')
            if fig:
                st.pyplot(fig)
        
        if 'platform_name' in df_filtrado.columns:
            fig = criar_grafico_pizza(df_filtrado, 'platform_name', 'Distribuição por Plataforma')
            if fig:
                st.pyplot(fig)
    
    with col2:
        if 'os_version' in df_filtrado.columns:
            fig = criar_grafico_barras(df_filtrado, 'os_version', 'Distribuição de Sistemas Operacionais')
            if fig:
                st.pyplot(fig)
        
        if 'policy_applied' in df_filtrado.columns:
            fig = criar_grafico_pizza(df_filtrado, 'policy_applied', 'Políticas Aplicadas')
            if fig:
                st.pyplot(fig)

    # Exportar Word
    st.markdown("---")
    st.markdown('<div class="section-header">📤 Exportação</div>', unsafe_allow_html=True)
    
    doc_buffer = export_word(df_filtrado, selected_company)
    
    st.download_button(
        label=f"📥 Exportar Relatório Word - {selected_company}",
        data=doc_buffer,
        file_name=f"{selected_company}_Relatorio_Executivo_{datetime.now().strftime('%Y%m%d_%H%M')}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        use_container_width=True
    )

# ----------------------------
# Upload CSV externo
# ----------------------------
st.markdown("---")
st.markdown('<div class="section-header">📤 Upload CSV Externo</div>', unsafe_allow_html=True)

uploaded_file = st.file_uploader("Envie seu arquivo CSV", type="csv", help="Faça upload de um arquivo CSV com dados compatíveis")
if uploaded_file:
    df_csv = pd.read_csv(uploaded_file)
    
    st.success(f"✅ {len(df_csv)} registros carregados do CSV!")
    
    # Aplicar os mesmos filtros
    df_csv_filtrado = aplicar_filtros(df_csv)
    
    # KPIs para CSV
    st.markdown('<div class="section-header">📊 KPIs - Dados CSV</div>', unsafe_allow_html=True)
    kpi_csv_cols = st.columns(4)
    
    with kpi_csv_cols[0]:
        st.markdown('<div class="kpi-card">', unsafe_allow_html=True)
        st.metric("Total Registros", len(df_csv_filtrado))
        st.markdown('</div>', unsafe_allow_html=True)
    
    # Gráficos para CSV
    st.markdown('<div class="section-header">📈 Gráficos - Dados CSV</div>', unsafe_allow_html=True)
    
    # Identificar colunas numéricas e categóricas automaticamente
    colunas_numericas = df_csv_filtrado.select_dtypes(include=[np.number]).columns.tolist()
    colunas_categoricas = df_csv_filtrado.select_dtypes(include=['object']).columns.tolist()
    
    col_csv1, col_csv2 = st.columns(2)
    
    with col_csv1:
        # Gráficos para colunas categóricas (máximo 2)
        for col in colunas_categoricas[:2]:
            fig = criar_grafico_barras(df_csv_filtrado, col, f'Distribuição - {col}')
            if fig:
                st.pyplot(fig)
    
    with col_csv2:
        # Gráficos para colunas numéricas (máximo 2)
        for col in colunas_numericas[:2]:
            if col in df_csv_filtrado.columns:
                fig, ax = plt.subplots(figsize=(4, 3))
                df_csv_filtrado[col].hist(bins=20, color='#D32F2F', alpha=0.8, ax=ax)
                ax.set_title(f'Distribuição - {col}', fontsize=11, fontweight='bold', color='white')
                ax.set_facecolor('#0E1117')
                fig.patch.set_facecolor('#0E1117')
                ax.tick_params(colors='white')
                ax.spines['bottom'].set_color('#3D3D5C')
                ax.spines['top'].set_color('#3D3D5C')
                ax.spines['right'].set_color('#3D3D5C')
                ax.spines['left'].set_color('#3D3D5C')
                plt.tight_layout()
                st.pyplot(fig)
    
    # Exportar dados CSV também
    doc_csv_buffer = export_word(df_csv_filtrado, "Dados CSV Externo")
    
    st.download_button(
        label="📥 Exportar Relatório Word - Dados CSV",
        data=doc_csv_buffer,
        file_name=f"CSV_Relatorio_Executivo_{datetime.now().strftime('%Y%m%d_%H%M')}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        use_container_width=True
    )
